# -*- coding: utf-8 -*-
"""
Crivo — servidor do produto (FastAPI). Serve a interface bonita (static/) e
expõe a API que faz a busca real e a expansão de descritores. A IA (triagem,
ficha, manuscrito) entra na etapa 2. Rodar:  python run.py
"""
from pathlib import Path

from fastapi import FastAPI
from fastapi.responses import FileResponse, Response
from pydantic import BaseModel

import motor_busca as mb
import motor_ia as mia
import motor_cita as mc

BASE = Path(__file__).parent
app = FastAPI(title="Crivo")


class BuscaReq(BaseModel):
    termo: str = ""
    bases: list[str] | None = None
    limite: int = 20
    usar_descritores: bool = True


class TriarReq(BaseModel):
    ref: dict
    criterios_inc: list[str] = []
    criterios_exc: list[str] = []
    pico: dict = {}


@app.get("/api/ia-status")
def ia_status():
    prov, modelo, ok, info = mia.provedor_ativo()
    return {"provedor": prov, "modelo": modelo, "ok": ok, "info": info}


@app.post("/api/triar")
def triar(req: TriarReq):
    pico = req.pico or {"desfecho": req.ref.get("_tema", "")}
    try:
        return mia.triar_referencia(req.ref, req.criterios_inc, req.criterios_exc, pico)
    except Exception as e:
        return {"decisao": "pendente", "motivo": f"erro IA: {e}", "criterio": ""}


class RefsReq(BaseModel):
    refs: list[dict] = []
    estilo: str = "vancouver"


@app.get("/api/estilos")
def estilos():
    return [{"id": e, "nome": mc.NOMES[e]} for e in mc.ESTILOS]


# Estruturas de artigo (o usuario escolhe; a IA escreve secao por secao, ancorada).
ESTRUTURAS = [
    {"id": "sintese", "nome": "Síntese rápida (1 seção)", "secoes": [
        {"titulo": "Análise crítica da evidência", "instrucao": ""}]},
    {"id": "sistematica", "nome": "Revisão Sistemática (PRISMA)", "secoes": [
        {"titulo": "Resumo", "instrucao": "Resumo estruturado e curto (contexto, objetivo, métodos, principais achados e conclusão), cerca de 200 palavras."},
        {"titulo": "Introdução", "instrucao": "Contextualize o tema e a lacuna do conhecimento; termine enunciando o objetivo/pergunta de pesquisa (PICO)."},
        {"titulo": "Métodos", "instrucao": "Escreva no estilo PRISMA: bases consultadas, estratégia de busca, critérios de inclusão/exclusão, processo de seleção e extração. É seção de metodologia — NÃO cite [n]."},
        {"titulo": "Resultados", "instrucao": "Apresente as características e os achados dos estudos incluídos, agrupando por temas; sustente cada afirmação com [n]."},
        {"titulo": "Discussão", "instrucao": "Interprete os achados, compare entre os estudos incluídos, aponte implicações para a prática e as limitações da evidência."},
        {"titulo": "Conclusão", "instrucao": "Conclusão objetiva, ancorada nos achados; sem trazer dados novos."}]},
    {"id": "narrativa", "nome": "Revisão Narrativa", "secoes": [
        {"titulo": "Introdução", "instrucao": "Apresente o tema, sua relevância e o objetivo do texto."},
        {"titulo": "Desenvolvimento", "instrucao": "Discuta o tema em blocos temáticos, integrando os estudos incluídos e citando por [n]."},
        {"titulo": "Considerações finais", "instrucao": "Feche com uma síntese reflexiva ancorada nos achados."}]},
    {"id": "integrativa", "nome": "Revisão Integrativa", "secoes": [
        {"titulo": "Introdução", "instrucao": "Contexto, justificativa e questão norteadora."},
        {"titulo": "Método", "instrucao": "Descreva as etapas da revisão integrativa (questão, busca nas bases, critérios, coleta). Seção de metodologia — NÃO cite [n]."},
        {"titulo": "Resultados", "instrucao": "Caracterize os estudos incluídos e agrupe os achados em categorias temáticas, citando [n]."},
        {"titulo": "Discussão", "instrucao": "Analise as categorias frente à literatura incluída, com implicações e limitações."},
        {"titulo": "Conclusão", "instrucao": "Síntese conclusiva ancorada nos achados."}]},
    {"id": "original", "nome": "Artigo Original", "secoes": [
        {"titulo": "Introdução", "instrucao": "Contexto, lacuna e objetivo/hipótese."},
        {"titulo": "Métodos", "instrucao": "Descreva o desenho, a fonte dos dados e a análise, com base nos estudos incluídos. Seção de metodologia — NÃO cite [n]."},
        {"titulo": "Resultados", "instrucao": "Apresente os achados dos estudos incluídos de forma objetiva, com [n]."},
        {"titulo": "Discussão", "instrucao": "Interprete, compare com a literatura incluída e aponte limitações."},
        {"titulo": "Conclusão", "instrucao": "Conclusão objetiva ancorada nos achados."}]},
]


@app.get("/api/estruturas")
def estruturas():
    return ESTRUTURAS


# Revistas-alvo: aplicam as normas da revista (estilo de ref., resumo, descritores).
# A maioria das revistas de enfermagem BR usa Vancouver (citacao numerica [n]),
# que casa com a ancoragem do Crivo. Verificar sempre as "Instrucoes aos autores"
# vigentes da revista antes de submeter.
_REGRA_BASE = ("Referencias no estilo Vancouver (citacao numerica [n]). "
               "Resumo estruturado em ~150 palavras com Objetivo, Metodo(s), "
               "Resultados e Conclusao. Incluir 3 a 5 descritores DeCS/MeSH ao "
               "final do resumo. Portugues cientifico, impessoal (evitar 1a pessoa). "
               "Titulo conciso. ")

REVISTAS = [
    {"id": "generico", "nome": "Genérico (sem revista específica)", "qualis": "", "estilo": "", "regras": ""},
    {"id": "reben", "nome": "Rev. Brasileira de Enfermagem (REBEn)", "qualis": "Qualis A1", "estilo": "vancouver",
     "regras": _REGRA_BASE + "Padrao REBEn/IMRAD; destacar contribuicoes para a enfermagem."},
    {"id": "texto_contexto", "nome": "Texto & Contexto Enfermagem", "qualis": "Qualis A1", "estilo": "vancouver",
     "regras": _REGRA_BASE + "Deixar clara a questao norteadora e o metodo de revisao."},
    {"id": "rlae", "nome": "Rev. Latino-Americana de Enfermagem (RLAE)", "qualis": "Qualis A1", "estilo": "vancouver",
     "regras": _REGRA_BASE + "Enfatizar implicacoes para a pratica de enfermagem."},
    {"id": "reeusp", "nome": "Rev. Esc. Enfermagem da USP (REEUSP)", "qualis": "Qualis A1", "estilo": "vancouver",
     "regras": _REGRA_BASE + "Estrutura IMRAD."},
    {"id": "acta", "nome": "Acta Paulista de Enfermagem", "qualis": "Qualis A1", "estilo": "vancouver",
     "regras": _REGRA_BASE + "Resumo com Objetivo, Metodos, Resultados, Conclusao."},
    {"id": "anna_nery", "nome": "Escola Anna Nery", "qualis": "Qualis A2", "estilo": "vancouver",
     "regras": _REGRA_BASE + "Enfase na relevancia para o cuidado."},
    {"id": "rgenf", "nome": "Rev. Gaúcha de Enfermagem", "qualis": "Qualis A2", "estilo": "vancouver",
     "regras": _REGRA_BASE},
    {"id": "cogitare", "nome": "Cogitare Enfermagem", "qualis": "Qualis A3", "estilo": "vancouver",
     "regras": _REGRA_BASE},
    {"id": "enf_foco", "nome": "Enfermagem em Foco (COFEN)", "qualis": "Qualis B1", "estilo": "vancouver",
     "regras": _REGRA_BASE},
    {"id": "rbso", "nome": "Rev. Bras. Saúde Ocupacional", "qualis": "Qualis B1", "estilo": "vancouver",
     "regras": _REGRA_BASE},
]


@app.get("/api/revistas")
def revistas():
    return REVISTAS


class TituloReq(BaseModel):
    refs: list[dict] = []
    tema: str = ""
    idioma: str = "pt"


@app.post("/api/titulo")
def titulo(req: TituloReq):
    prov, _m, ok, _i = mia.provedor_ativo()
    if not ok:
        return {"titulo": req.tema}
    try:
        return {"titulo": mia.gerar_titulo(req.tema, req.refs, idioma=req.idioma)}
    except Exception:
        return {"titulo": req.tema}


class DocxReq(BaseModel):
    titulo: str = ""
    texto: str = ""
    figura_png: str = ""      # PNG em base64 (dataURL) do fluxograma PRISMA
    figura_legenda: str = ""


def _par_negrito(doc, texto):
    """Adiciona um paragrafo convertendo **negrito** em runs em negrito."""
    p = doc.add_paragraph()
    for i, seg in enumerate(texto.split("**")):
        if not seg:
            continue
        run = p.add_run(seg)
        if i % 2 == 1:
            run.bold = True
    return p


@app.post("/api/docx")
def docx(req: DocxReq):
    """Converte o artigo montado (markdown simples) num arquivo Word .docx."""
    import io as _io
    from docx import Document

    doc = Document()
    for raw in (req.texto or "").split("\n"):
        s = raw.strip()
        if s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=1)
        elif s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=0)
        elif len(s) > 2 and s.startswith("*") and s.endswith("*"):
            p = doc.add_paragraph()
            r = p.add_run(s.strip("*"))
            r.italic = True
        elif not s:
            doc.add_paragraph("")
        else:
            _par_negrito(doc, raw.rstrip())

    # figura PRISMA embutida (se enviada)
    if req.figura_png:
        try:
            import base64
            from docx.shared import Inches
            b64 = req.figura_png.split(",", 1)[-1]
            img = _io.BytesIO(base64.b64decode(b64))
            doc.add_heading(req.figura_legenda or "Figura 1 — Fluxograma PRISMA 2020", level=1)
            doc.add_picture(img, width=Inches(6.0))
        except Exception:
            pass

    buf = _io.BytesIO()
    doc.save(buf)
    nome = "".join(c for c in (req.titulo or "artigo")[:60] if c.isalnum() or c in " -_") or "artigo"
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{nome}.docx"'},
    )


@app.post("/api/referencias")
def referencias(req: RefsReq):
    itens = [mc.formatar(r, req.estilo) for r in req.refs]
    return {"estilo": req.estilo, "n": len(itens), "itens": itens,
            "texto": mc.formatar_lista(req.refs, req.estilo)}


class RascunhoReq(BaseModel):
    refs: list[dict] = []
    tema: str = ""
    secao: str = "Análise crítica da evidência"
    instrucoes: str = ""
    idioma: str = "pt"


@app.post("/api/rascunhar")
def rascunhar(req: RascunhoReq):
    """Escreve uma seção do trabalho ANCORADA nos artigos incluídos (só cita o
    que foi lido; nada inventado). Precisa da IA ligada."""
    prov, _mod, ok, _info = mia.provedor_ativo()
    if not ok:
        return {"secao": req.secao, "texto": "",
                "erro": "IA desligada — configure a chave (GOOGLE_API_KEY ou GROQ_API_KEY)."}
    pico = {"desfecho": req.tema}
    try:
        texto = mia.rascunhar_secao(req.secao, pico, req.refs,
                                    texto_extra=req.instrucoes, idioma=req.idioma)
        return {"secao": req.secao, "texto": texto}
    except Exception as e:
        return {"secao": req.secao, "texto": "", "erro": str(e)}


@app.get("/api/bases")
def bases():
    return [{"id": k, "nome": v} for k, v in mb.NOMES_BASES.items()]


@app.post("/api/expandir")
def expandir(req: BuscaReq):
    return mb.expandir_descritores(req.termo)


@app.post("/api/buscar")
def buscar(req: BuscaReq):
    termo = (req.termo or "").strip()
    if not termo:
        return {"registros": [], "por_base": {}, "erros": {}, "descritores": {}}

    desc = mb.expandir_descritores(termo) if req.usar_descritores else {"mesh": []}
    mesh = desc.get("mesh", [])[:4]
    # query enriquecida: termo livre OR descritores MeSH (entre aspas)
    if mesh:
        termos = [termo] + [f'"{m}"' for m in mesh]
        query = "(" + " OR ".join(dict.fromkeys(termos)) + ")"
    else:
        query = termo

    bases = req.bases or [b for b in mb.NOMES_BASES if b != "lilacs"]
    alvo = {b: query for b in bases if b in mb.BASES}
    res = mb.buscar_todas(alvo, limite=req.limite)
    res["descritores"] = desc
    res["query"] = query
    return res


def _autores(r):
    return r.get("autores_est") or r.get("autores") or []


def gerar_ris(refs):
    """Formato RIS (Zotero, Mendeley, EndNote)."""
    out = []
    for r in refs:
        out.append("TY  - JOUR")
        out.append("TI  - " + (r.get("titulo") or ""))
        for a in _autores(r):
            out.append("AU  - " + str(a))
        if r.get("ano"):
            out.append("PY  - " + str(r["ano"]))
        if r.get("revista"):
            out.append("JO  - " + str(r["revista"]))
        if r.get("volume"):
            out.append("VL  - " + str(r["volume"]))
        if r.get("numero"):
            out.append("IS  - " + str(r["numero"]))
        if r.get("paginas"):
            out.append("SP  - " + str(r["paginas"]))
        if r.get("doi"):
            out.append("DO  - " + str(r["doi"]))
        if r.get("url"):
            out.append("UR  - " + str(r["url"]))
        if r.get("resumo"):
            out.append("AB  - " + str(r["resumo"]))
        out.append("ER  - ")
        out.append("")
    return "\n".join(out)


def gerar_bibtex(refs):
    """Formato BibTeX (LaTeX, Zotero, Mendeley)."""
    blocos = []
    for i, r in enumerate(refs, 1):
        aut = _autores(r)
        base = (aut[0].split(",")[0].split()[0] if aut else "ref")
        key = ("".join(c for c in base if c.isalnum()) or "ref") + str(r.get("ano") or "") + str(i)
        f = ["  title = {" + (r.get("titulo") or "") + "}"]
        if aut:
            f.append("  author = {" + " and ".join(str(a) for a in aut) + "}")
        if r.get("ano"):
            f.append("  year = {" + str(r["ano"]) + "}")
        if r.get("revista"):
            f.append("  journal = {" + str(r["revista"]) + "}")
        if r.get("volume"):
            f.append("  volume = {" + str(r["volume"]) + "}")
        if r.get("numero"):
            f.append("  number = {" + str(r["numero"]) + "}")
        if r.get("paginas"):
            f.append("  pages = {" + str(r["paginas"]) + "}")
        if r.get("doi"):
            f.append("  doi = {" + str(r["doi"]) + "}")
        blocos.append("@article{" + key + ",\n" + ",\n".join(f) + "\n}")
    return "\n\n".join(blocos)


class ExportReq(BaseModel):
    refs: list[dict] = []
    formato: str = "ris"


@app.post("/api/exportar")
def exportar(req: ExportReq):
    fmt = (req.formato or "ris").lower()
    if fmt == "bibtex":
        return {"formato": "bibtex", "ext": "bib", "texto": gerar_bibtex(req.refs)}
    return {"formato": "ris", "ext": "ris", "texto": gerar_ris(req.refs)}


@app.get("/")
def index():
    for p in (BASE / "index.html", BASE / "static" / "index.html"):
        if p.exists():
            return FileResponse(str(p))
    return {"erro": "index.html nao encontrado"}
